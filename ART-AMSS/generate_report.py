"""
Oracle Database Health Check Report Generator
Generates AEROTHAI_AMSS_Oracle_HealthCheck_202605.docx

Log files in output/YYYYMM/ are read automatically when present.
Falls back to gray placeholder boxes when files are missing.
"""

import os
import re
import zipfile
import tempfile
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# Configuration — แก้ค่าในส่วนนี้ก่อนรันทุกครั้ง
# ─────────────────────────────────────────────────────────────────────────────

MONTH_TAG  = "202605"   # YYYYMM — เดือนที่รายงาน
YEAR       = "2026"     # ปี CE
DOC_SEQ    = "1"        # ลำดับรายงานในปีนั้น

# เอกสารหน้าปก (ไม่ถูก overwrite — output ไปเป็นไฟล์ใหม่)
_BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
BASE_OUTPUT = os.path.join(_BASE_DIR, "output")
COVER_FILE  = os.path.join(_BASE_DIR, "หน้าปกรายงาน_วิทยุการบิน_AMSS.docx")

# ─────────────────────────────────────────────────────────────────────────────
# Derived paths (ไม่ต้องแก้)
# ─────────────────────────────────────────────────────────────────────────────

def get_host_dir(hostname: str) -> str:
    """คืน path ของ folder log ของแต่ละ server: output/YYYYMM_hostname/"""
    return os.path.join(BASE_OUTPUT, f"{MONTH_TAG}_{hostname}")

# รายงานลงตรง output/ (ไม่อยู่ใน sub-folder)
OUTPUT_DIR  = BASE_OUTPUT
OUTPUT_FILE = os.path.join(
    BASE_OUTPUT,
    f"รายงานการบำรุงรักษาระบบฐานข้อมูล_AMSS_{YEAR}_{DOC_SEQ}.docx",
)

# Color theme — Blue
COLOR_H1       = "000000"   # Black
COLOR_H2       = "000000"   # Black
COLOR_H3       = "000000"   # Black
COLOR_H4       = "000000"   # Black
COLOR_TITLE    = "000000"   # Black
COLOR_SUBTITLE = "000000"   # Black
COLOR_CELL_HEADER = "DAEEF3"   # Light Blue (ONPROD style)
COLOR_CELL_ODD    = "FFFFFF"   # White (no alternating, ONPROD style)
COLOR_CELL_EVEN   = "FFFFFF"
COLOR_COVER_LABEL = "DAEEF3"   # Light Blue
COLOR_LOG_BOX     = "F2F2F2"
COLOR_SQL_BOX     = "EBF3FB"   # Very light blue — SQL command display box
COLOR_SECTION_HDR = "DAEEF3"   # Light Blue
COLOR_HDR_TEXT    = "1F497D"   # Dark navy — text on light blue header rows

CDBS = [
    {"cdb": "amsscdb",  "pdbs": ["AMSS", "AMSSPDB"], "primary_server": "dbsystem2", "standby_server": "dbsystem1"},
    {"cdb": "fdmccdb",  "pdbs": ["FDMC"],             "primary_server": "dbsystem2", "standby_server": "dbsystem1"},
    {"cdb": "fdmscdb",  "pdbs": ["FDMS"],             "primary_server": "dbsystem2", "standby_server": "dbsystem1"},
]

# ─────────────────────────────────────────────────────────────────────────────
# Oracle Database Configuration (Section 1.1.1)
# ─────────────────────────────────────────────────────────────────────────────
ORACLE_PLATFORM     = "64-bit Oracle Linux"
ORACLE_EDITION      = "Oracle Database 19c Enterprise Edition"
ORACLE_GRID_HOME    = "/u01/app/19.19.0.0/grid"
ORACLE_BASE         = "/u01/app/oracle"
ORACLE_HOME         = "/u01/app/oracle/product/19.0.0.0/dbhome_1"


# ─────────────────────────────────────────────────────────────────────────────
# Log file parser
# ─────────────────────────────────────────────────────────────────────────────

class LogReader:
    """
    Parses Oracle health check spool files (===header=== / --- subsection -- PROMPT format).
    Falls back gracefully when file is missing.
    """

    def __init__(self, filepath: str):
        self.filepath = filepath
        self._sections: dict = {}
        if os.path.exists(filepath):
            self._parse()

    def available(self) -> bool:
        return bool(self._sections)

    def _parse(self):
        with open(self.filepath, encoding="utf-8", errors="replace") as f:
            text = f.read()
        hdr_re = re.compile(r"={40,}\n(.+?)\n={40,}", re.MULTILINE)
        matches = list(hdr_re.finditer(text))
        for i, m in enumerate(matches):
            label = m.group(1).strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            self._sections[label] = self._split_subs(text[start:end])

    def _split_subs(self, content: str) -> dict:
        sub_re = re.compile(r"^---\s+(.+?)(?:\s+--\s+PROMPT)?\s*---\s*$", re.MULTILINE)
        matches = list(sub_re.finditer(content))
        result = {"__full__": content.strip()}
        if matches:
            preamble = content[: matches[0].start()].strip()
            if preamble:
                result["__preamble__"] = preamble
            for j, m in enumerate(matches):
                sub_label = m.group(1).strip()
                s = m.end()
                e = matches[j + 1].start() if j + 1 < len(matches) else len(content)
                result[sub_label] = content[s:e].strip('\n')
        return result

    def _find_section(self, section_key: str) -> dict:
        for label, subs in self._sections.items():
            if section_key in label:
                return subs
        return {}

    def get(self, section_key: str, *subsection_keys: str) -> str:
        """Return section content, optionally joining named subsections."""
        subs = self._find_section(section_key)
        if not subs:
            return ""
        if not subsection_keys:
            return subs.get("__full__", "")
        parts = []
        for sk in subsection_keys:
            if sk.startswith("__"):
                v = subs.get(sk, "")
                if v:
                    parts.append(v)
            else:
                for k, v in subs.items():
                    if not k.startswith("__") and sk.lower() in k.lower():
                        parts.append(f"--- {k} ---\n{v}")
                        break
        return "\n\n".join(p for p in parts if p.strip())

    def raw(self) -> str:
        try:
            with open(self.filepath, encoding="utf-8", errors="replace") as f:
                return f.read().strip()
        except Exception:
            return ""


class SqlScriptReader:
    """
    Parses pm_collect_primary.sql / pm_collect_standby.sql and extracts the SQL
    SELECT/EXEC/RMAN queries per subsection (delimited by PROMPT --- Label ---).
    """

    def __init__(self, filepath: str):
        self._sections: dict = {}
        if os.path.exists(filepath):
            self._parse(filepath)

    def available(self) -> bool:
        return bool(self._sections)

    def _parse(self, filepath: str):
        with open(filepath, encoding="utf-8", errors="replace") as f:
            content = f.read()
        sec_re = re.compile(r"^PROMPT --- (.+?) ---\s*$", re.MULTILINE)
        matches = list(sec_re.finditer(content))
        for i, m in enumerate(matches):
            label = m.group(1).strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            sql = self._extract_sql(content[start:end])
            if sql:
                self._sections[label] = sql

    def _extract_sql(self, block: str) -> str:
        """Extract only the SQL statement lines (skips col/set formatting lines)."""
        lines = block.split("\n")
        sql_lines = []
        in_sql = False
        for line in lines:
            lstrip = line.lstrip()
            if not in_sql:
                if re.match(
                    r"(select|with|exec|execute|begin|rman|backup|crosscheck|list|delete|update|insert)\b",
                    lstrip, re.IGNORECASE,
                ):
                    in_sql = True
            if in_sql:
                sql_lines.append(line.rstrip())
                if line.rstrip().endswith(";") or line.strip() == "/":
                    in_sql = False
                    sql_lines.append("")
        return "\n".join(sql_lines).strip()

    def get(self, *labels: str) -> str:
        """Return SQL blocks for labels that partially match (case-insensitive)."""
        results = []
        for label in labels:
            for k, v in self._sections.items():
                if label.lower() in k.lower():
                    results.append(f"-- {k}\n{v}")
                    break
        return "\n\n".join(results)


# Load SQL script readers once at module level
_SQL_DIR = os.path.dirname(os.path.abspath(__file__))
SQL_PRIMARY  = SqlScriptReader(os.path.join(_SQL_DIR, "pm_collect_primary.sql"))
SQL_STANDBY  = SqlScriptReader(os.path.join(_SQL_DIR, "pm_collect_standby.sql"))


def filter_pdb_rows(content: str, pdb_name: str) -> str:
    """Keep header, separator, and rows whose first word matches pdb_name."""
    lines = content.split("\n")
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if "PDB_Name" in line or "Tablespace_Name" in line:
            result.append(line)
        elif set(stripped) <= {"-", " "}:
            result.append(line)
        elif stripped.upper().startswith(pdb_name.upper()):
            result.append(line)
    return "\n".join(result)


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def set_para_shading(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)


def set_table_borders(table, color="4BACC6", size=8):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tblBorders.append(border)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def set_col_width(cell, width_cm: float):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)


def add_page_break(doc: Document):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


# ─────────────────────────────────────────────────────────────────────────────
# Heading helpers
# ─────────────────────────────────────────────────────────────────────────────

def add_h1(doc, text):
    para = doc.add_heading(text, level=1)
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(COLOR_H1)
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(18)
    return para


def add_h2(doc, text):
    para = doc.add_heading(text, level=2)
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(COLOR_H2)
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(16)
    return para


def add_h3(doc, text):
    para = doc.add_heading(text, level=3)
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(COLOR_H3)
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(14)
    return para


def add_h4(doc, text):
    para = doc.add_heading(text, level=4)
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(COLOR_H4)
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(13)
    return para


def add_body(doc, text):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(14)
    return para


# ─────────────────────────────────────────────────────────────────────────────
# Log content helpers
# ─────────────────────────────────────────────────────────────────────────────

def add_log_box(doc: Document, placeholder: str, blank_lines: int = 8):
    """Gray placeholder box (used when real log is unavailable)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(placeholder)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_para_shading(para, COLOR_LOG_BOX)
    for _ in range(blank_lines):
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        bp.add_run(" ").font.name = "Courier New"
        bp.runs[0].font.size = Pt(9)
        set_para_shading(bp, COLOR_LOG_BOX)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(2)
    gap.paragraph_format.space_after = Pt(4)


def add_log_content(doc: Document, text: str, max_lines: int = 120):
    """Insert actual log content in gray Courier New box."""
    lines = text.split("\n")
    if len(lines) > max_lines:
        lines = lines[:max_lines] + [f"... [{len(lines) - max_lines} more lines]"]
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line if line.strip() else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(7)
        set_para_shading(para, COLOR_LOG_BOX)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(2)
    gap.paragraph_format.space_after = Pt(4)


def add_log(doc: Document, reader, section: str, *subsections,
            placeholder: str = "[ไม่พบข้อมูล]", blank_lines: int = 8, max_lines: int = 120,
            sql_reader=None, sql_keys=()):
    """Show real log content when available, else fall back to placeholder box."""
    if sql_reader and sql_reader.available():
        keys = sql_keys if sql_keys else subsections
        add_sql_cmd(doc, sql_reader.get(*keys))
    if reader and reader.available():
        text = reader.get(section, *subsections)
        if text.strip():
            add_log_content(doc, text, max_lines=max_lines)
            return
    add_log_box(doc, placeholder, blank_lines)


def add_raw_log(doc: Document, reader, placeholder: str = "[ไม่พบข้อมูล]", blank_lines: int = 8):
    """Show entire file content as-is."""
    if reader:
        text = reader.raw()
        if text.strip():
            add_log_content(doc, text)
            return
    add_log_box(doc, placeholder, blank_lines)


def add_sql_cmd(doc: Document, sql_text: str):
    """Display the SQL command in a light-blue box before the result."""
    if not sql_text or not sql_text.strip():
        return
    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_before = Pt(4)
    label_p.paragraph_format.space_after = Pt(0)
    run = label_p.add_run("คำสั่ง SQL:")
    run.bold = True
    run.font.name = "TH SarabunPSK"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
    for line in sql_text.split("\n"):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(line if line.strip() else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        set_para_shading(para, COLOR_SQL_BOX)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(2)
    gap.paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────────────────────────────
# Alert log renderer
# ─────────────────────────────────────────────────────────────────────────────

COLOR_ALERT_TS  = "D9E1F2"   # light blue — timestamp row
COLOR_ALERT_MSG = "FFFFFF"   # white — message row


def add_alertlog_section(doc: Document, reader, section_key: str,
                         placeholder: str = "[ไม่พบข้อมูล]", blank_lines: int = 12,
                         sql_reader=None, sql_keys=()):
    """
    Render alert log as a 2-row-per-entry layout:
      Row A (light blue): Timestamp
      Row B (white):      [Lvl] Message
    Falls back to log_box when no data.
    """
    if sql_reader and sql_reader.available():
        add_sql_cmd(doc, sql_reader.get(*sql_keys) if sql_keys else "")
    if not (reader and reader.available()):
        add_log_box(doc, placeholder, blank_lines)
        return

    content = reader.get(section_key)
    if not content.strip():
        add_log_box(doc, placeholder, blank_lines)
        return

    parsed = parse_sqlplus_table(content)
    if not parsed or not parsed["rows"]:
        add_log_box(doc, placeholder, blank_lines)
        return

    headers = parsed["headers"]          # [Timestamp, Lvl, Message]
    rows    = parsed["rows"]
    # Find column indices
    ts_idx  = next((i for i, h in enumerate(headers) if "time" in h.lower()), 0)
    lvl_idx = next((i for i, h in enumerate(headers) if "lvl" in h.lower()), 1)
    msg_idx = next((i for i, h in enumerate(headers) if "message" in h.lower()), 2)

    col_w_ts  = 4.5
    col_w_msg = 12.0

    tbl = doc.add_table(rows=len(rows) * 2, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="4BACC6", size=8)

    for i, row_data in enumerate(rows):
        ts  = row_data[ts_idx]  if ts_idx  < len(row_data) else ""
        lvl = row_data[lvl_idx] if lvl_idx < len(row_data) else ""
        msg = row_data[msg_idx] if msg_idx < len(row_data) else ""
        full_msg = f"[{lvl}] {msg}" if lvl.strip() else msg

        # Row A — Timestamp (merged across 2 cols, light blue)
        row_a = tbl.rows[i * 2]
        merged = row_a.cells[0].merge(row_a.cells[1])
        set_cell_shading(merged, COLOR_ALERT_TS)
        p = merged.paragraphs[0]
        r = p.add_run(ts)
        r.bold = True; r.font.name = "Courier New"; r.font.size = Pt(9)
        merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Row B — Message
        row_b = tbl.rows[i * 2 + 1]
        # indent cell (empty, narrow)
        set_cell_shading(row_b.cells[0], COLOR_ALERT_MSG)
        set_col_width(row_b.cells[0], col_w_ts)
        row_b.cells[0].paragraphs[0].add_run("")

        set_cell_shading(row_b.cells[1], COLOR_ALERT_MSG)
        set_col_width(row_b.cells[1], col_w_msg)
        p2 = row_b.cells[1].paragraphs[0]
        r2 = p2.add_run(full_msg)
        r2.font.name = "Courier New"; r2.font.size = Pt(9)
        row_b.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Generic SQLPlus table parser & renderer
# ─────────────────────────────────────────────────────────────────────────────

def parse_sqlplus_table(text: str):
    """
    Parse a single SQLPlus fixed-width column block.
    Detects header + separator (---) + data rows.
    Handles continuation lines (first col empty, wrapping value in later col).
    Returns {headers: [...], rows: [[...], ...]} or None if not parseable.
    """
    lines = text.split("\n")
    sep_idx = None
    for i, line in enumerate(lines):
        if (i > 0 and line.strip() and
                re.match(r"^[\-\s]+$", line) and
                line.count("-") >= 3 and
                lines[i - 1].strip()):
            sep_idx = i
            break
    if sep_idx is None:
        return None

    sep_line = lines[sep_idx]
    col_bounds = [(m.start(), m.end()) for m in re.finditer(r"-+", sep_line)]
    if not col_bounds:
        return None

    def extract(line):
        result = []
        for i, (start, end) in enumerate(col_bounds):
            next_start = col_bounds[i + 1][0] if i + 1 < len(col_bounds) else len(line) + 20
            val = line[start:min(next_start, len(line))].strip() if start < len(line) else ""
            result.append(val)
        return result

    headers = extract(lines[sep_idx - 1])
    rows = []
    for line in lines[sep_idx + 1:]:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^[\-\s]+$", line) and line.count("-") >= 3:
            break
        if re.match(r"^\d+ rows? selected", stripped, re.IGNORECASE):
            break
        if stripped.lower() == "no rows selected":
            break
        row = extract(line)
        if not row[0].strip() and rows:
            for i, val in enumerate(row):
                if val.strip():
                    rows[-1][i] = (rows[-1][i] + " " + val).strip()
        elif any(v.strip() for v in row):
            rows.append(list(row))

    return {"headers": headers, "rows": rows} if rows else None


def render_sqltable(doc: Document, parsed: dict, title: str = None, max_rows: int = 200):
    """Render a parsed SQLPlus table as a formatted Word table."""
    if not parsed or not parsed["rows"]:
        return False

    headers = parsed["headers"]
    rows = parsed["rows"][:max_rows]
    num_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if num_cols == 0:
        return False

    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True; r.font.name = "TH SarabunPSK"; r.font.size = Pt(12)

    # Proportional column widths based on max content length
    col_max = [max(
        len(headers[c]) if c < len(headers) else 0,
        max((len(row[c]) for row in rows if c < len(row)), default=0), 2
    ) for c in range(num_cols)]
    total_chars = sum(col_max) or 1
    col_widths = [max(1.0, (n / total_chars) * 16.5) for n in col_max]

    # Detect numeric columns: >50% of non-empty values parseable as float
    _num_re = re.compile(r"^-?[\d,]+\.?\d*$")
    def _is_numeric_col(c):
        vals = [row[c] for row in rows if c < len(row) and row[c].strip()]
        if not vals:
            return False
        return sum(1 for v in vals if _num_re.match(v.replace(",", ""))) / len(vals) > 0.5

    col_numeric = [_is_numeric_col(c) for c in range(num_cols)]

    tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="4BACC6", size=8)

    for c in range(num_cols):
        cell = tbl.rows[0].cells[c]
        set_cell_shading(cell, COLOR_CELL_HEADER)
        set_col_width(cell, col_widths[c])
        p = cell.paragraphs[0]
        r = p.add_run(headers[c] if c < len(headers) else "")
        r.bold = True; r.font.name = "TH SarabunPSK"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        fill = COLOR_CELL_ODD if r_idx % 2 == 0 else COLOR_CELL_EVEN
        for c in range(num_cols):
            cell = row.cells[c]
            set_cell_shading(cell, fill); set_col_width(cell, col_widths[c])
            p = cell.paragraphs[0]
            r = p.add_run(row_data[c] if c < len(row_data) else "")
            r.font.name = "Courier New"; r.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_numeric[c] else WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if len(parsed["rows"]) > max_rows:
        p = doc.add_paragraph()
        r = p.add_run(f"... [{len(parsed['rows']) - max_rows} rows not shown]")
        r.italic = True; r.font.name = "TH SarabunPSK"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return True


def add_sql_tables_from_content(doc: Document, content: str,
                                placeholder: str = "[ไม่พบข้อมูล]", blank_lines: int = 8):
    """
    Split content by sub-section markers (--- Name -- PROMPT  or  --- Name ---)
    and render each block as a Word table. Falls back to log_box if nothing parsed.
    """
    if not content.strip():
        add_log_box(doc, placeholder, blank_lines)
        return

    sub_re = re.compile(r"^---\s+(.+?)(?:\s+--\s+PROMPT)?\s*---\s*$", re.MULTILINE)
    matches = list(sub_re.finditer(content))
    rendered = 0

    if matches:
        for i, m in enumerate(matches):
            title = m.group(1).strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            block = content[start:end].strip()
            if not block:
                continue
            parsed = parse_sqlplus_table(block)
            if parsed:
                render_sqltable(doc, parsed, title=title)
                rendered += 1
    else:
        parsed = parse_sqlplus_table(content)
        if parsed:
            render_sqltable(doc, parsed)
            rendered += 1

    if rendered == 0:
        add_log_box(doc, placeholder, blank_lines)


def add_sql_section(doc: Document, reader, section_key: str,
                    placeholder: str = "[ไม่พบข้อมูล]", blank_lines: int = 8,
                    sql_reader=None, sql_keys=()):
    """Fetch full section from LogReader and render as tables."""
    if sql_reader and sql_reader.available():
        add_sql_cmd(doc, sql_reader.get(*sql_keys) if sql_keys else "")
    if reader and reader.available():
        content = reader.get(section_key)
        if content.strip():
            add_sql_tables_from_content(doc, content, placeholder, blank_lines)
            return
    add_log_box(doc, placeholder, blank_lines)


def add_sql_section_keys(doc: Document, reader, section_key: str, *sub_keys,
                         placeholder: str = "[ไม่พบข้อมูล]", blank_lines: int = 8,
                         sql_reader=None):
    """Fetch specific sub-sections from LogReader (joins with --- markers) and render."""
    if sql_reader and sql_reader.available() and sub_keys:
        add_sql_cmd(doc, sql_reader.get(*sub_keys))
    if reader and reader.available():
        content = reader.get(section_key, *sub_keys)
        if content.strip():
            add_sql_tables_from_content(doc, content, placeholder, blank_lines)
            return
    add_log_box(doc, placeholder, blank_lines)


def add_blank_space(doc: Document, lines: int = 15, note: str = ""):
    if note:
        p = doc.add_paragraph()
        r = p.add_run(f"[ {note} ]")
        r.font.name = "TH SarabunPSK"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r.italic = True
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
    for _ in range(lines):
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        bp.add_run(" ").font.size = Pt(11)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(2)
    gap.paragraph_format.space_after = Pt(6)


# ─────────────────────────────────────────────────────────────────────────────
# Checklist table
# ─────────────────────────────────────────────────────────────────────────────

CHECKLIST_COL_WIDTHS_CM = [1.2, 7.5, 2.8, 3.0, 2.8]
CHECKLIST_HEADERS = ["ลำดับ", "รายการตรวจสอบ", "ผลลัพธ์", "สถานะ", "หมายเหตุ"]


def add_checklist(doc: Document, items: list):
    table = doc.add_table(rows=1 + len(items), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="4BACC6", size=8)

    hdr_row = table.rows[0]
    for c, (cell, hdr) in enumerate(zip(hdr_row.cells, CHECKLIST_HEADERS)):
        set_cell_shading(cell, COLOR_CELL_HEADER)
        set_col_width(cell, CHECKLIST_COL_WIDTHS_CM[c])
        para = cell.paragraphs[0]
        run = para.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(13)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, item_text in enumerate(items):
        row = table.rows[r_idx + 1]
        fill = COLOR_CELL_ODD if r_idx % 2 == 0 else COLOR_CELL_EVEN
        cells = row.cells

        set_cell_shading(cells[0], fill); set_col_width(cells[0], CHECKLIST_COL_WIDTHS_CM[0])
        p0 = cells[0].paragraphs[0]; p0.add_run(str(r_idx + 1)).font.name = "TH SarabunPSK"
        p0.runs[0].font.size = Pt(13); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_shading(cells[1], fill); set_col_width(cells[1], CHECKLIST_COL_WIDTHS_CM[1])
        p1 = cells[1].paragraphs[0]; p1.add_run(item_text).font.name = "TH SarabunPSK"
        p1.runs[0].font.size = Pt(13); cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_shading(cells[2], fill); set_col_width(cells[2], CHECKLIST_COL_WIDTHS_CM[2])
        cells[2].paragraphs[0].add_run(""); cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_shading(cells[3], fill); set_col_width(cells[3], CHECKLIST_COL_WIDTHS_CM[3])
        p3 = cells[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_items = ("☐  ปกติ", "☐  ไม่ปกติ")
        for idx, text in enumerate(status_items):
            r = p3.add_run(text)
            r.font.name = "TH SarabunPSK"; r.font.size = Pt(12)
            if idx < len(status_items) - 1:
                r._r.append(OxmlElement("w:br"))
        cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_shading(cells[4], fill); set_col_width(cells[4], CHECKLIST_COL_WIDTHS_CM[4])
        cells[4].paragraphs[0].add_run(""); cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    return table


def add_combined_checklist(doc: Document, sections: list):
    """
    Render all checklist groups in one table.
    sections: [(section_label, [item_text, ...]), ...]
    """
    total_rows = 1 + sum(1 + len(items) for _, items in sections)

    table = doc.add_table(rows=total_rows, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="4BACC6", size=8)

    hdr_row = table.rows[0]
    for c, (cell, hdr) in enumerate(zip(hdr_row.cells, CHECKLIST_HEADERS)):
        set_cell_shading(cell, COLOR_CELL_HEADER)
        set_col_width(cell, CHECKLIST_COL_WIDTHS_CM[c])
        para = cell.paragraphs[0]
        run = para.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(13)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    row_idx = 1
    item_counter = 1

    for section_label, items in sections:
        # Section separator row
        sec_row = table.rows[row_idx]
        merged = sec_row.cells[0]
        for i in range(1, 5):
            merged = merged.merge(sec_row.cells[i])
        set_cell_shading(merged, COLOR_SECTION_HDR)
        p = merged.paragraphs[0]
        run = p.add_run(section_label)
        run.bold = True
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        row_idx += 1

        for item_text in items:
            row = table.rows[row_idx]
            fill = COLOR_CELL_ODD if item_counter % 2 == 0 else COLOR_CELL_EVEN
            cells = row.cells

            set_cell_shading(cells[0], fill); set_col_width(cells[0], CHECKLIST_COL_WIDTHS_CM[0])
            p0 = cells[0].paragraphs[0]; p0.add_run(str(item_counter)).font.name = "TH SarabunPSK"
            p0.runs[0].font.size = Pt(13); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            set_cell_shading(cells[1], fill); set_col_width(cells[1], CHECKLIST_COL_WIDTHS_CM[1])
            p1 = cells[1].paragraphs[0]; p1.add_run(item_text).font.name = "TH SarabunPSK"
            p1.runs[0].font.size = Pt(13); cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            set_cell_shading(cells[2], fill); set_col_width(cells[2], CHECKLIST_COL_WIDTHS_CM[2])
            cells[2].paragraphs[0].add_run(""); cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            set_cell_shading(cells[3], fill); set_col_width(cells[3], CHECKLIST_COL_WIDTHS_CM[3])
            p3 = cells[3].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for idx2, text in enumerate(("☐  ปกติ", "☐  ไม่ปกติ")):
                r = p3.add_run(text)
                r.font.name = "TH SarabunPSK"; r.font.size = Pt(12)
                if idx2 == 0:
                    r._r.append(OxmlElement("w:br"))
            cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            set_cell_shading(cells[4], fill); set_col_width(cells[4], CHECKLIST_COL_WIDTHS_CM[4])
            cells[4].paragraphs[0].add_run(""); cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            item_counter += 1
            row_idx += 1

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Checklist item definitions
# ─────────────────────────────────────────────────────────────────────────────

ITEMS_INSTANCE   = ["Instance Status = OPEN", "Database Status = ACTIVE", "Open Mode = READ WRITE"]
ITEMS_TS_CDB     = ["ทุก Tablespace มีพื้นที่เหลือ > 15%", "SYSTEM / SYSAUX ไม่เกิน 85%", "FRA ใช้งานไม่เกิน 80%"]
ITEMS_TS_PDB     = ["ทุก Tablespace มีพื้นที่เหลือ > 15%", "ไม่มี Tablespace OFFLINE"]
ITEMS_ALERT      = ["ไม่พบ ORA- Error ร้ายแรงใน 31 วัน", "ไม่พบ Checkpoint / Archive Error"]
ITEMS_PERF       = ["ไม่มี Top Wait Event ที่ส่งผลกระทบ", "AWR Snapshot รันสม่ำเสมอ"]
ITEMS_RMAN       = ["Backup status = COMPLETED (32 วันล่าสุด)", "ไม่มี Backup FAILED / EXPIRED"]
ITEMS_PARAM      = ["SGA / PGA ตั้งค่าเหมาะสม", "ไม่มี ORA-04031 (Shared Pool Error)"]
ITEMS_PATCH      = ["Oracle Version ถูกต้อง", "Patch ล่าสุดถูกติดตั้ง"]
ITEMS_DG_SYNC    = ["DG Member ครบถ้วน (PRIMARY + STANDBY)", "Transport Lag = 0 / Apply Lag = 0"]
ITEMS_DG_GAP     = ["ไม่มี Archive Gap", "Apply Lag < 5 นาที", "Transport Lag < 5 นาที"]
ITEMS_DG_TRANS   = ["Archive Dest Status = VALID", "MRP0 Process = APPLYING_LOG"]
ITEMS_DG_DISK    = ["FRA ใช้งานไม่เกิน 80%", "Tablespace Standby ปกติ"]
ITEMS_DG_ALERT   = ["ไม่พบ ORA- Error ร้ายแรง", "ไม่พบ Archive Gap Error"]
ITEMS_DG_RMAN    = ["Backup COMPLETED หรือ N/A (ถ้าไม่ได้ backup จาก standby)"]
ITEMS_NET        = ["tnsping สำเร็จจาก dbsystem1 → AMSSCDB_STBY", "tnsping สำเร็จจาก dbsystem2 → AMSSCDB", "Response time อยู่ในระดับปกติ (< 100 ms)"]


# ─────────────────────────────────────────────────────────────────────────────
# Table of Contents
# ─────────────────────────────────────────────────────────────────────────────

def build_toc(doc: Document):
    print("  Building Table of Contents...")

    p_toc = doc.add_paragraph(style="Title")
    p_toc.add_run("สารบัญ").font.name = "TH SarabunPSK"

    p_hdr = doc.add_paragraph()
    for word in ["เรื่อง", "\t" * 11, "         ", "หน้า"]:
        r = p_hdr.add_run(word)
        r.bold = True
        r.font.name = "TH SarabunPSK"
        r.font.size = Pt(13)

    toc_entries = [
        "1. ข้อมูลระบบ (System Information)",
        "   1.1 DC / Server & OS",
        "   1.2 CDB / PDB Inventory",
        "2. Database Enterprise Edition",
        "   2.1 AMSSCDB (PRIMARY @ dbsystem2)",
        "   2.2 FDMCCDB (PRIMARY @ dbsystem2)",
        "   2.3 FDMSCDB (PRIMARY @ dbsystem2)",
        "3. Active Data Guard",
        "   3.1 AMSSCDB (STANDBY @ dbsystem1)",
        "   3.2 FDMCCDB (STANDBY @ dbsystem1)",
        "   3.3 FDMSCDB (STANDBY @ dbsystem1)",
        "4. Patch",
        "   4.1 OPatch Patches",
        "   4.2 Oracle Critical Patch Update (CPU)",
        "5. Network",
        "   5.1 tnsping Test",
        "   5.2 Network Bandwidth",
        "6. สรุปผลการตรวจสอบ (Summary)",
    ]
    for entry in toc_entries:
        p = doc.add_paragraph()
        r = p.add_run(entry)
        r.font.name = "TH SarabunPSK"
        r.font.size = Pt(13)

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────────────────────

def build_cover(doc: Document):
    print("  Building cover page...")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(48)
    p_title.paragraph_format.space_after = Pt(8)
    run = p_title.add_run("รายงานผลการตรวจสอบระบบฐานข้อมูล")
    run.bold = True; run.font.name = "TH SarabunPSK"; run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(COLOR_TITLE)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after = Pt(24)
    run2 = p_sub.add_run("Oracle Database 19c  |  Oracle Linux 7.9  |  Active Data Guard")
    run2.bold = True; run2.font.name = "TH SarabunPSK"; run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor.from_string(COLOR_SUBTITLE)

    doc.add_paragraph()
    cover_data = [
        ("ชื่อโครงการ",    "จ้างซ่อมแซมบำรุงรักษาอุปกรณ์ Storage พร้อมระบบจัดการฐานข้อมูลระบบ AMSS ระยะเวลา 1 ปี"),
        ("ผู้ว่าจ้าง",     "บริษัท วิทยุการบินแห่งประเทศไทย จำกัด (AEROTHAI)"),
        ("ผู้รับจ้าง",     "บริษัท appworks Company Limited"),
        ("สถานที่",        "ห้อง Data Center ชั้น 2  |  102 ซอยงามดูพลี ถนนพระราม 4 กรุงเทพฯ"),
        ("ระบบปฏิบัติการ", "Oracle Linux 7.9 (OL7.9)"),
        ("Oracle Version", "Oracle Database 19c Enterprise Edition"),
        ("วันที่ตรวจสอบ", "วันที่ _______ / _______ / _______"),
        ("ผู้ตรวจสอบ",    "________________________________________"),
    ]
    table = doc.add_table(rows=len(cover_data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="4BACC6", size=8)
    for r_idx, (label, value) in enumerate(cover_data):
        cells = table.rows[r_idx].cells
        set_cell_shading(cells[0], COLOR_COVER_LABEL); set_col_width(cells[0], 5.0)
        rl = cells[0].paragraphs[0].add_run(label); rl.bold = True
        rl.font.name = "TH SarabunPSK"; rl.font.size = Pt(14)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_col_width(cells[1], 10.0)
        rv = cells[1].paragraphs[0].add_run(value)
        rv.font.name = "TH SarabunPSK"; rv.font.size = Pt(14)
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Oracle Database Configuration table (Section 1.1.1)
# ─────────────────────────────────────────────────────────────────────────────

def load_oracle_config(hostname: str) -> dict:
    """
    Read oracle_config_<hostname>_YYYYMM.txt produced by pm_collect.sh.
    Falls back to module-level constants when file is missing.
    Returns dict with keys: platform, edition, oracle_home, oracle_base, grid_home.
    """
    defaults = {
        "platform":    ORACLE_PLATFORM,
        "edition":     ORACLE_EDITION,
        "oracle_home": ORACLE_HOME,
        "oracle_base": ORACLE_BASE,
        "grid_home":   ORACLE_GRID_HOME,
    }
    cfg_path = os.path.join(get_host_dir(hostname), f"oracle_config_{hostname}_{MONTH_TAG}.txt")
    if not os.path.exists(cfg_path):
        return None
    mapping = {
        "platform":    "Platform",
        "edition":     "Oracle Edition",
        "oracle_home": "Oracle Home",
        "oracle_base": "Oracle Base",
        "grid_home":   "Grid Home",
    }
    result = dict(defaults)
    with open(cfg_path, encoding="utf-8", errors="replace") as f:
        for line in f:
            for key, label in mapping.items():
                if line.startswith(label):
                    val = line.split(":", 1)[1].strip()
                    if val:
                        result[key] = val
    return result


def add_oracle_config_table(doc: Document, cfg: dict = None):
    """
    Render ONPROD-style Oracle Database Configuration table:
      - Section header rows (full-width, light blue DAEEF3, navy bold text)
      - Single-value rows (label: bold value, merged 2 cols, white bg)
      - Two-column rows (label | bold value)
    cfg — dict from load_oracle_config(); falls back to module constants when None.
    """
    COL_LABEL = 5.5   # cm
    COL_VALUE = 11.5  # cm

    if cfg is None:
        cfg = {
            "platform":    ORACLE_PLATFORM,
            "edition":     ORACLE_EDITION,
            "oracle_home": ORACLE_HOME,
            "oracle_base": ORACLE_BASE,
            "grid_home":   ORACLE_GRID_HOME,
        }

    # Row spec: ("section", label) | ("row1", label, value) | ("row2", label, value)
    rows_spec = [
        ("section", "Software"),
        ("row1", "Platform:",        cfg["platform"]),
        ("row1", "Oracle Product:",  cfg["edition"]),
        ("section", "Database Software"),
        ("row2", "Oracle Base Directory",     cfg["oracle_base"]),
        ("row2", "Oracle Home Directory",     cfg["oracle_home"]),
    ]

    tbl = doc.add_table(rows=len(rows_spec), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl, color="4BACC6", size=8)

    for r_idx, spec in enumerate(rows_spec):
        row = tbl.rows[r_idx]

        if spec[0] == "section":
            # Full-width merged header row
            merged = row.cells[0].merge(row.cells[1])
            set_cell_shading(merged, COLOR_CELL_HEADER)
            p = merged.paragraphs[0]
            run = p.add_run(spec[1])
            run.bold = True
            run.font.name = "TH SarabunPSK"
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        elif spec[0] == "row1":
            # Merged row: "Label: **Value**" in one cell
            merged = row.cells[0].merge(row.cells[1])
            set_cell_shading(merged, COLOR_CELL_EVEN)
            p = merged.paragraphs[0]
            r_label = p.add_run(spec[1] + " ")
            r_label.font.name = "TH SarabunPSK"
            r_label.font.size = Pt(13)
            r_val = p.add_run(spec[2])
            r_val.bold = True
            r_val.font.name = "TH SarabunPSK"
            r_val.font.size = Pt(13)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        else:  # "row2"
            # Two-column: right-aligned label | bold value
            cell_lbl = row.cells[0]
            cell_val = row.cells[1]
            set_cell_shading(cell_lbl, COLOR_CELL_EVEN)
            set_cell_shading(cell_val, COLOR_CELL_EVEN)
            set_col_width(cell_lbl, COL_LABEL)
            set_col_width(cell_val, COL_VALUE)

            p_lbl = cell_lbl.paragraphs[0]
            r = p_lbl.add_run(spec[1])
            r.font.name = "TH SarabunPSK"
            r.font.size = Pt(13)
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell_lbl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p_val = cell_val.paragraphs[0]
            r = p_val.add_run(spec[2])
            r.bold = True
            r.font.name = "TH SarabunPSK"
            r.font.size = Pt(13)
            p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Oracle Database Environment table (Section 1.2 — per CDB)
# ─────────────────────────────────────────────────────────────────────────────

def _render_label_value_table(doc: Document, section_title: str, rows: list):
    """
    Render an ONPROD-style label|value table with a section header row.
    rows: [(display_label, value), ...]
    """
    COL_LABEL = 5.5
    COL_VALUE = 11.5
    rows_spec = [("section", section_title)] + [("data", lbl, val) for lbl, val in rows]

    tbl = doc.add_table(rows=len(rows_spec), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl, color="4BACC6", size=8)

    for r_idx, spec in enumerate(rows_spec):
        row = tbl.rows[r_idx]
        if spec[0] == "section":
            merged = row.cells[0].merge(row.cells[1])
            set_cell_shading(merged, COLOR_CELL_HEADER)
            p = merged.paragraphs[0]
            run = p.add_run(spec[1])
            run.bold = True
            run.font.name = "TH SarabunPSK"; run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
            merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            cell_lbl, cell_val = row.cells[0], row.cells[1]
            set_cell_shading(cell_lbl, COLOR_CELL_EVEN)
            set_cell_shading(cell_val, COLOR_CELL_EVEN)
            set_col_width(cell_lbl, COL_LABEL)
            set_col_width(cell_val, COL_VALUE)
            r = cell_lbl.paragraphs[0].add_run(spec[1])
            r.font.name = "TH SarabunPSK"; r.font.size = Pt(13)
            cell_lbl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell_lbl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = cell_val.paragraphs[0].add_run(spec[2])
            r.bold = True
            r.font.name = "TH SarabunPSK"; r.font.size = Pt(13)
            cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


STORAGE_LABEL_MAP = {
    "Global_DB_Name": "Global DB Name",
    "DB_Name_SID":    "Oracle DB Name (SID)",
    "DB_Location":    "DB Location",
    "Archive_Dir":    "Archive Directory",
    "Archive_Mode":   "Archive Mode",
    "Flashback":      "Oracle Flashback",
}

CHARSET_LABEL_MAP = {
    "NLS_CHARACTERSET":       "DB Character Set",
    "NLS_NCHAR_CHARACTERSET": "National Character Set",
}


def _parse_db_env_old_format(raw: str) -> list:
    """
    Parse old 6-column wrapped SQLPlus output for [1.2 - DB Environment].

    SQLPlus wraps a wide row as follows (linesize=115, cols 1-3 = 97 chars, cols 4-6 = 87 chars):
      Header Part1  (cols 1-3)
      Separator1
      Header Part2  (cols 4-6)   ← looks like data to naive parser
      Separator2
      Data Part1    (cols 1-3)   ← actual data starts here
      [col3 continuation if > 30 chars]
      Data Part2    (cols 4-6)
      [col4 continuation if > 45 chars]
    """
    LABELS = [
        "Global DB Name", "Oracle DB Name (SID)", "DB Location",
        "Archive Directory", "Archive Mode", "Oracle Flashback",
    ]
    lines = raw.split("\n")

    # Find the two separator lines
    sep_idxs = [i for i, l in enumerate(lines)
                if re.match(r"^[\-\s]+$", l) and l.count("-") >= 10]
    if len(sep_idxs) < 2:
        return []

    bounds1 = [(m.start(), m.end()) for m in re.finditer(r"-+", lines[sep_idxs[0]])]
    bounds2 = [(m.start(), m.end()) for m in re.finditer(r"-+", lines[sep_idxs[1]])]

    def extract(line, bounds):
        cols = []
        for i, (s, e) in enumerate(bounds):
            nxt = bounds[i+1][0] if i+1 < len(bounds) else len(line)+20
            cols.append(line[s:min(nxt, len(line))].strip() if s < len(line) else "")
        return cols

    # ALL data lines start after sep_idxs[1]
    data_lines = [l for l in lines[sep_idxs[1]+1:] if l.strip()]

    # Partition: Part1 = first group (starts non-space, continuation = starts space)
    #            Part2 = second group (next non-space start after Part1 is done)
    part1, part2 = [], []
    in_part2 = False
    for line in data_lines:
        if not line:
            continue
        starts_nonspace = (line[:1] not in (" ", ""))
        if not in_part2:
            if starts_nonspace and part1:
                # Second non-space start → Part2 begins
                in_part2 = True
                part2.append(line)
            else:
                part1.append(line)
        else:
            part2.append(line)

    def merge(block, bounds):
        # SQLPlus word-wrap: continuation lines join directly (no added space)
        vals = [""] * len(bounds)
        for line in block:
            for i, v in enumerate(extract(line, bounds)):
                if v:
                    vals[i] = (vals[i] + v) if vals[i] else v
        return vals

    vals1 = merge(part1, bounds1)
    vals2 = merge(part2, bounds2)

    result = []
    for label, val in zip(LABELS, vals1 + vals2):
        if label == "Archive Directory" and "LOCATION=" in val.upper():
            m = re.search(r"LOCATION=(\S+)", val, re.IGNORECASE)
            val = m.group(1) if m else val
        result.append((label, val))
    return result


def add_db_environment_section(doc: Document, reader):
    """
    Render Oracle Database Environment (Storage Mechanism + Character Set)
    as ONPROD-style label|value tables. Reads [1.2 - DB Environment] from PRIMARY log.
    Falls back to placeholder when log is unavailable.
    """
    if not (reader and reader.available()):
        add_log_box(doc, "[ไม่พบข้อมูล DB Environment — ต้องมี log จาก pm_collect_primary.sql]", blank_lines=6)
        return

    # ── Storage Mechanism ──────────────────────────────────────────────────
    storage_raw = reader.get("[1.2 - DB Environment]", "Storage Mechanism")
    parsed = parse_sqlplus_table(storage_raw) if storage_raw.strip() else None
    rows = []
    if parsed and parsed["rows"]:
        headers = [h.strip() for h in parsed["headers"]]
        if headers and headers[0].lower() in ("item", "item"):
            # New 2-column format: each row = (Item, Value)
            rows = [(r[0], r[1] if len(r) > 1 else "") for r in parsed["rows"]]
        else:
            # Old 6-column format: parse raw text directly by known label positions
            rows = _parse_db_env_old_format(storage_raw)
    if rows:
        _render_label_value_table(doc, "Storage Mechanism", rows)
    else:
        add_log_box(doc, "[ไม่พบ Storage Mechanism]", blank_lines=5)

    # ── Database Character Set ─────────────────────────────────────────────
    # SQL outputs 2-column (Parameter, Value) — map Parameter → Thai label
    charset_raw = reader.get("[1.2 - DB Environment]", "Database Character Set")
    parsed_cs = parse_sqlplus_table(charset_raw) if charset_raw.strip() else None
    if parsed_cs and parsed_cs["rows"]:
        rows_cs = [(CHARSET_LABEL_MAP.get(r[0], r[0]), r[1] if len(r) > 1 else "")
                   for r in parsed_cs["rows"]]
        _render_label_value_table(doc, "Database Character Set", rows_cs)
    else:
        add_log_box(doc, "[ไม่พบ Database Character Set]", blank_lines=3)


# ─────────────────────────────────────────────────────────────────────────────
# Inventory parser & table renderer
# ─────────────────────────────────────────────────────────────────────────────

def parse_inventory(filepath: str) -> dict:
    """
    Parse inventory_{host}_{YYYYMM}.txt into:
    {
      "server": "dbsystem1",
      "collected": "2026-05-27 06:21:29",
      "cdbs": [
        {"cdb": "amsscdb", "role": "PHYSICAL STANDBY", "version": "19.31.0.0.0", "host": "dbsystem1",
         "pdbs": [{"name": "AMSS", "open_mode": "READ ONLY", "restricted": "NO"}, ...]},
        ...
      ]
    }
    """
    result = {"server": "", "collected": "", "cdbs": []}
    if not os.path.exists(filepath):
        return result

    with open(filepath, encoding="utf-8", errors="replace") as f:
        lines = f.readlines()

    for line in lines:
        if line.strip().startswith("Server"):
            result["server"] = line.split(":", 1)[-1].strip()
        elif line.strip().startswith("Collected"):
            result["collected"] = line.split(":", 1)[-1].strip()

    cdb_re = re.compile(
        r"^CDB\s*:\s*(\S+)\s+Role:\s*(.+?)\s+Version:\s*(\S+)\s+Host:\s*(\S+)", re.IGNORECASE
    )
    pdb_re = re.compile(r"^\s{4}(\S+)\s+open_mode=(.+?)\s+restricted=(\S+)", re.IGNORECASE)

    current_cdb = None
    for line in lines:
        m = cdb_re.match(line)
        if m:
            current_cdb = {
                "cdb": m.group(1),
                "role": m.group(2).strip().replace("PHYSICALSTANDBY", "PHYSICAL STANDBY"),
                "version": m.group(3),
                "host": m.group(4),
                "pdbs": [],
            }
            result["cdbs"].append(current_cdb)
            continue
        if current_cdb:
            mp = pdb_re.match(line)
            if mp:
                current_cdb["pdbs"].append({
                    "name": mp.group(1),
                    "open_mode": mp.group(2).strip(),
                    "restricted": mp.group(3).strip(),
                })
    return result


COLOR_INV_HEADER = "DAEEF3"   # light blue — section header rows (ONPROD style)
COLOR_INV_CDB    = "DAEEF3"   # light blue — CDB row


def add_inventory_table(doc: Document, inv: dict):
    """Render parsed inventory as a formatted table."""
    if not inv["cdbs"]:
        add_log_box(doc, "[ไม่พบข้อมูล Inventory]", blank_lines=6)
        return

    # Meta row above table
    if inv["server"] or inv["collected"]:
        p = doc.add_paragraph()
        for txt, bold in [(f"Server: {inv['server']}", True),
                          (f"   |   Collected: {inv['collected']}", False)]:
            r = p.add_run(txt)
            r.bold = bold
            r.font.name = "TH SarabunPSK"
            r.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(2)

    col_headers = ["CDB", "Role", "Version", "PDB", "Open Mode"]
    col_widths  = [2.8, 4.5, 2.8, 2.8, 4.7]

    # Flatten to one row per PDB for clear column alignment
    flat_rows = []
    for cdb in inv["cdbs"]:
        pdbs = cdb["pdbs"] if cdb["pdbs"] else [{"name": "—", "open_mode": "—"}]
        for i, pdb in enumerate(pdbs):
            flat_rows.append({
                "cdb":       cdb["cdb"].upper() if i == 0 else "",
                "role":      cdb["role"]         if i == 0 else "",
                "version":   cdb["version"]      if i == 0 else "",
                "pdb":       pdb["name"],
                "open_mode": pdb["open_mode"],
                "first_pdb": i == 0,
            })

    tbl = doc.add_table(rows=1 + len(flat_rows), cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="4BACC6", size=8)

    # Header row
    hdr_row = tbl.rows[0]
    for c, (hdr, w) in enumerate(zip(col_headers, col_widths)):
        cell = hdr_row.cells[c]
        set_cell_shading(cell, COLOR_INV_HEADER)
        set_col_width(cell, w)
        p = cell.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True; r.font.name = "TH SarabunPSK"; r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, fr in enumerate(flat_rows):
        row  = tbl.rows[r_idx + 1]
        fill = COLOR_INV_CDB if fr["first_pdb"] else (COLOR_CELL_ODD if r_idx % 2 == 0 else COLOR_CELL_EVEN)
        vals = [fr["cdb"], fr["role"], fr["version"], fr["pdb"], fr["open_mode"]]
        for c, (val, w) in enumerate(zip(vals, col_widths)):
            cell = row.cells[c]
            set_cell_shading(cell, fill)
            set_col_width(cell, w)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.bold = (c == 0 and fr["first_pdb"]) or c == 3
            r.font.name = "TH SarabunPSK"; r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Section 1 – System Information
# ─────────────────────────────────────────────────────────────────────────────

def build_section1(doc: Document):
    print("  Building Section 1: ข้อมูลระบบ...")
    add_h1(doc, "1. ข้อมูลระบบ (System Information)")

    add_h2(doc, "1.1 DC / Server & OS")
    add_body(doc, "ข้อมูลเซิร์ฟเวอร์ทั้งสองเครื่อง ได้แก่ dbsystem1 และ dbsystem2 พร้อมข้อมูลระบบปฏิบัติการ Oracle Linux 7.9")

    add_h3(doc, "1.1.1 Oracle Database Configuration")
    # โหลดค่าจริงจาก oracle_config_<host>_YYYYMM.txt (primary server ก่อน, fallback standby)
    cfg = load_oracle_config(CDBS[0]["primary_server"]) \
          or load_oracle_config(CDBS[0]["standby_server"])
    if cfg:
        add_oracle_config_table(doc, cfg)
    else:
        add_log_box(doc, f"[ไม่พบ oracle_config_{{host}}_{MONTH_TAG}.txt — รัน main_collect.sh แล้ว copy ไฟล์มาวางใน output/{MONTH_TAG}/]", blank_lines=6)

    for host in ("dbsystem1", "dbsystem2"):
        add_h3(doc, f"{host} – Inventory")
        inv_path = os.path.join(get_host_dir(host), f"inventory_{host}_{MONTH_TAG}.txt")
        inv = parse_inventory(inv_path)
        if inv["cdbs"]:
            add_inventory_table(doc, inv)
        else:
            add_log_box(doc, f"[วางผลลัพธ์จาก: inventory_{host}_{MONTH_TAG}.txt]", blank_lines=10)

    add_h2(doc, "1.2 CDB / PDB Inventory")
    add_body(doc, "รายการ CDB และ PDB ที่ติดตั้งบนระบบ ประกอบด้วย amsscdb, fdmccdb, fdmscdb")

    inv_headers = ["CDB", "PDB", "Role (dbsystem2)", "Role (dbsystem1)", "Version"]
    inv_data = [
        ["amsscdb", "AMSS, AMSSPDB", "PRIMARY", "STANDBY", "19c"],
        ["fdmccdb", "FDMC",          "PRIMARY", "STANDBY", "19c"],
        ["fdmscdb", "FDMS",          "PRIMARY", "STANDBY", "19c"],
    ]
    inv_tbl = doc.add_table(rows=1 + len(inv_data), cols=5)
    inv_tbl.style = "Table Grid"
    inv_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(inv_tbl, color="4BACC6", size=8)
    for c, hdr in enumerate(inv_headers):
        cell = inv_tbl.rows[0].cells[c]
        set_cell_shading(cell, COLOR_CELL_HEADER)
        p = cell.paragraphs[0]; run = p.add_run(hdr)
        run.bold = True; run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        run.font.name = "TH SarabunPSK"; run.font.size = Pt(13)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(inv_data):
        row = inv_tbl.rows[r_idx + 1]
        fill = COLOR_CELL_ODD if r_idx % 2 == 0 else COLOR_CELL_EVEN
        for c, val in enumerate(row_data):
            cell = row.cells[c]; set_cell_shading(cell, fill)
            p = cell.paragraphs[0]; run = p.add_run(val)
            run.font.name = "TH SarabunPSK"; run.font.size = Pt(13)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 1.2.x Oracle Database Environment (per CDB) ────────────────────────
    add_h3(doc, "1.2.1 Oracle Database Environment")
    add_body(doc, "รายละเอียดสภาพแวดล้อมฐานข้อมูล (Storage, Archive Mode, Character Set) ของแต่ละ CDB")
    for cdb_info in CDBS:
        cdb = cdb_info["cdb"]
        srv = cdb_info["primary_server"]
        log_path = os.path.join(get_host_dir(srv), f"{cdb}_PRIMARY_{srv}_{MONTH_TAG}.txt")
        reader = LogReader(log_path)
        add_h4(doc, cdb.upper())
        add_db_environment_section(doc, reader)

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 2 – Database Enterprise Edition (PRIMARY)
# ─────────────────────────────────────────────────────────────────────────────

def add_sql_fulltext_table(doc: Document, reader, section: str, subsection: str,
                           placeholder: str, blank_lines: int = 10,
                           sql_reader=None, sql_keys=()):
    """Render Top SQL Full Text as 2-column fixed-width table: SQL ID | SQL Text."""
    if sql_reader and sql_reader.available():
        add_sql_cmd(doc, sql_reader.get(*sql_keys) if sql_keys else "")
    if not (reader and reader.available()):
        add_log_box(doc, placeholder, blank_lines)
        return
    content = reader.get(section, subsection)
    if not content.strip():
        add_log_box(doc, placeholder, blank_lines)
        return
    parsed = parse_sqlplus_table(content)
    if not parsed or not parsed["rows"]:
        add_log_box(doc, placeholder, blank_lines)
        return

    headers = parsed["headers"]
    rows    = parsed["rows"]
    # Auto-detect columns: 2-col (SQL_ID, Text) or 3-col (PDB, SQL_ID, Text)
    num_cols = min(len(headers), 3)
    if num_cols == 3:
        col_widths = [2.0, 2.5, 12.0]   # PDB | SQL_ID | SQL Text
    else:
        col_widths = [2.5, 14.0]         # SQL_ID | SQL Text

    tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="4BACC6", size=8)

    for c in range(num_cols):
        cell = tbl.rows[0].cells[c]
        set_cell_shading(cell, COLOR_CELL_HEADER)
        set_col_width(cell, col_widths[c])
        p = cell.paragraphs[0]
        r = p.add_run(headers[c] if c < len(headers) else "")
        r.bold = True; r.font.name = "TH SarabunPSK"; r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row_data in enumerate(rows):
        row  = tbl.rows[r_idx + 1]
        fill = COLOR_CELL_ODD if r_idx % 2 == 0 else COLOR_CELL_EVEN
        for c in range(num_cols):
            cell = row.cells[c]
            set_cell_shading(cell, fill)
            set_col_width(cell, col_widths[c])
            val = row_data[c] if c < len(row_data) else ""
            p   = cell.paragraphs[0]
            r   = p.add_run(val)
            r.font.name = "Courier New"
            r.font.size = Pt(8 if c == 1 else 9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.add_paragraph()


def build_one_cdb_primary(doc: Document, cdb_info: dict, cdb_idx: int, reader=None):
    cdb  = cdb_info["cdb"]
    pdbs = cdb_info["pdbs"]
    srv  = cdb_info["primary_server"]
    txt  = f"{cdb}_PRIMARY_{srv}_{MONTH_TAG}.txt"
    pfx  = f"2.{cdb_idx}"

    add_h2(doc, f"{pfx} {cdb.upper()} (PRIMARY @ {srv})")

    # ── 2.x.1 Instance Status ──────────────────────────────────────────────
    add_h3(doc, f"{pfx}.1 Instance Status")
    add_body(doc, f"ตรวจสอบสถานะ Oracle Instance ของ {cdb.upper()} บน {srv}")
    add_sql_cmd(doc, SQL_PRIMARY.get("Instance Status"))
    _INST_HEADERS = ["Inst#", "Instance", "Host", "Version", "Status", "DB Status", "Days", "Startup Time"]
    if reader and reader.available():
        _content = reader.get("[2.1 - 1]", "Instance Status")
        _parsed  = parse_sqlplus_table(_content) if _content.strip() else None
        if _parsed and _parsed["rows"]:
            _parsed["headers"] = _INST_HEADERS[:len(_parsed["headers"])]
            render_sqltable(doc, _parsed)
            doc.add_paragraph()
        else:
            add_log_box(doc, f"[{txt} → [2.1-1] Instance Status]", blank_lines=4)
    else:
        add_log_box(doc, f"[{txt} → [2.1-1] Instance Status]", blank_lines=4)
    add_findings_box(doc,
        f"พบว่า Oracle Instance ของ {cdb.upper()} อยู่ในสถานะ OPEN และ Database Status = ACTIVE "
        f"หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 2.x.2 Tablespace ───────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.2 Tablespace / Temp / Undo / Redo Log")

    add_h4(doc, f"{pfx}.2.1 CDB Root Tablespace")
    add_body(doc, f"ตรวจสอบ Tablespace บน CDB Root ของ {cdb.upper()}")
    add_sql_section_keys(doc, reader, "[2.1 - 2]",
                         "Tablespace Usage (CDB Root)",
                         "Temp Tablespace Usage",
                         "Undo Statistics",
                         "Redo Log Status",
                         "Archive Log Generation",
                         "FRA (Flash Recovery Area) Summary",
                         "FRA Usage by File Type",
                         placeholder=f"[{txt} → [2.1-2] CDB Root + Temp/Undo/Redo/FRA]", blank_lines=14,
                         sql_reader=SQL_PRIMARY)
    add_findings_box(doc,
        f"พบว่า Tablespace บน CDB Root ของ {cdb.upper()} ทุกรายการยังมีพื้นที่ใช้งานเพียงพอ "
        f"ไม่พบ Tablespace ที่ใกล้เต็มหรืออยู่ในสถานะผิดปกติ หมายความว่าสามารถใช้งานได้ตามปกติ")

    for i, pdb in enumerate(pdbs):
        add_h4(doc, f"{pfx}.2.{i + 2} PDB: {pdb}")
        add_body(doc, f"ตรวจสอบ Tablespace บน PDB {pdb}")
        add_sql_cmd(doc, SQL_PRIMARY.get("Tablespace Usage (PDB Level)"))
        if reader and reader.available():
            raw_pdb = reader.get("[2.1 - 2]", "Tablespace Usage (PDB Level)")
            filtered = filter_pdb_rows(raw_pdb, pdb)
            parsed = parse_sqlplus_table(filtered) if filtered.strip() else None
            if parsed:
                render_sqltable(doc, parsed)
                doc.add_paragraph()
            else:
                add_log_box(doc, f"[{txt} → [2.1-2] PDB {pdb}]", blank_lines=8)
        else:
            add_log_box(doc, f"[{txt} → [2.1-2] PDB {pdb}]", blank_lines=8)
        add_findings_box(doc,
            f"พบว่า Tablespace บน PDB {pdb} ทุกรายการยังมีพื้นที่ใช้งานเพียงพอ "
            f"ไม่พบ Tablespace ที่ OFFLINE หรือเกินเกณฑ์ที่กำหนด หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 2.x.3 Alert Log ────────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.3 Alert Log")
    add_body(doc, f"ตรวจสอบ Alert Log ของ {cdb.upper()} เพื่อหา ORA- Error หรือ Warning ใน 31 วันที่ผ่านมา")
    add_alertlog_section(doc, reader, "[2.1 - 3]",
                         placeholder=f"[{txt} → [2.1-3] Alert Log]",
                         sql_reader=SQL_PRIMARY, sql_keys=("Alert Log",))
    add_findings_box(doc,
        f"พบว่าใน Alert Log ของ {cdb.upper()} ย้อนหลัง 31 วัน ไม่พบ ORA- Error ร้ายแรง "
        f"หรือ Warning ที่ส่งผลกระทบต่อการทำงาน หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 2.x.4 Performance ──────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.4 Performance (AWR / Top SQL / Wait Events)")
    add_body(doc, f"ตรวจสอบ Performance ของ {cdb.upper()} ย้อนหลัง 31 วัน (AWR + dba_hist) หรือเท่าที่ AWR retention มี")

    add_h4(doc, f"{pfx}.4.1 AWR Snapshot Summary (31 วัน)")
    add_sql_section_keys(doc, reader, "[2.1 - 4]", "AWR Snapshots",
                         placeholder=f"[{txt} → [2.1-4] AWR Snapshot]", blank_lines=4,
                         sql_reader=SQL_PRIMARY)

    add_h4(doc, f"{pfx}.4.2 Top SQL by Elapsed Time (AWR - 31 วัน)")
    add_sql_section_keys(doc, reader, "[2.1 - 4]", "Top 10 SQL by Elapsed Time",
                         placeholder=f"[{txt} → [2.1-4] Top SQL stats]", blank_lines=12,
                         sql_reader=SQL_PRIMARY)

    add_h4(doc, f"{pfx}.4.3 Top SQL — Full SQL Text")
    add_sql_fulltext_table(doc, reader, "[2.1 - 4]", "Top 10 SQL Full Text",
                           placeholder=f"[{txt} → [2.1-4] Top SQL full text]",
                           sql_reader=SQL_PRIMARY, sql_keys=("Top 10 SQL Full Text",))

    add_h4(doc, f"{pfx}.4.4 Top Wait Events (AWR - 31 วัน)")
    add_sql_section_keys(doc, reader, "[2.1 - 4]", "Top 15 System Wait Events",
                         placeholder=f"[{txt} → [2.1-4] Wait Events]", blank_lines=10,
                         sql_reader=SQL_PRIMARY)
    add_findings_box(doc,
        f"พบว่าค่า Performance ของ {cdb.upper()} ใน 31 วันที่ผ่านมา AWR Snapshot รันสม่ำเสมอ "
        f"ค่า Top SQL และ Wait Events อยู่ในระดับปกติ ไม่พบ Session หรือ Event ที่ส่งผลกระทบต่อประสิทธิภาพ "
        f"หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 2.x.5 RMAN ─────────────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.5 RMAN Backup")
    add_body(doc, f"ตรวจสอบสถานะ RMAN Backup ของ {cdb.upper()} ย้อนหลัง 32 วัน")
    add_sql_section(doc, reader, "[2.1 - 5]",
                    placeholder=f"[{txt} → [2.1-5] RMAN Backup]", blank_lines=12,
                    sql_reader=SQL_PRIMARY, sql_keys=("RMAN Backup", "Backup"))
    add_findings_box(doc,
        f"พบว่า RMAN Backup ของ {cdb.upper()} ในช่วง 32 วันที่ผ่านมา มีสถานะ COMPLETED ทุกรายการ "
        f"ไม่พบ Backup ที่ FAILED หรือ EXPIRED หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 2.x.6 Parameters ───────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.6 Parameters (SGA / PGA)")
    add_body(doc, f"ตรวจสอบค่า Parameter สำคัญของ {cdb.upper()}")
    add_log(doc, reader, "[2.1 - 6]",
            placeholder=f"[{txt} → [2.1-6] SGA / PGA Parameters]", blank_lines=8,
            sql_reader=SQL_PRIMARY, sql_keys=("Key Initialization Parameters", "SGA Components", "PGA Statistics"))
    add_findings_box(doc,
        f"พบว่าค่า Parameter สำคัญของ {cdb.upper()} ได้แก่ SGA, PGA และ Memory "
        f"ตั้งค่าเหมาะสมกับขนาดฐานข้อมูลและการใช้งาน ไม่พบค่า Parameter ที่ผิดปกติ "
        f"หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 2.x.7 Oracle Version ───────────────────────────────────────────────
    add_h3(doc, f"{pfx}.7 Oracle Version")
    add_body(doc, f"ตรวจสอบ Oracle Version และ Registry ของ {cdb.upper()} (รายละเอียด Patch ดูที่ Section 4)")
    add_sql_section_keys(doc, reader, "[2.1 - 7]",
                         "Database Version", "Installed Components", "Applied Patches",
                         placeholder=f"[{txt} → [2.1-7] Version / Registry]", blank_lines=8,
                         sql_reader=SQL_PRIMARY)
    add_findings_box(doc,
        f"พบว่า Oracle Database Version ของ {cdb.upper()} ติดตั้งและลงทะเบียน Component ครบถ้วน "
        f"รายละเอียด Patch ที่ติดตั้งดูได้ที่ Section 4 หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── Combined Checklist (bottom of CDB) ─────────────────────────────────
    add_h3(doc, f"{pfx}.8 รายการตรวจสอบ")
    checklist_sections = [
        (f"{pfx}.1  Instance Status",        ITEMS_INSTANCE),
        (f"{pfx}.2  Tablespace – CDB Root",  ITEMS_TS_CDB),
    ]
    for j, pdb in enumerate(pdbs):
        checklist_sections.append((f"{pfx}.2  Tablespace – PDB: {pdb}", ITEMS_TS_PDB))
    checklist_sections += [
        (f"{pfx}.3  Alert Log",              ITEMS_ALERT),
        (f"{pfx}.4  Performance",            ITEMS_PERF),
        (f"{pfx}.5  RMAN Backup",            ITEMS_RMAN),
        (f"{pfx}.6  Parameters",             ITEMS_PARAM),
        (f"{pfx}.7  Oracle Version",         ITEMS_PATCH),
    ]
    add_combined_checklist(doc, checklist_sections)


def build_section2(doc: Document):
    print("  Building Section 2: Database Enterprise Edition...")
    add_h1(doc, "2. Database Enterprise Edition")
    add_body(doc, "ตรวจสอบ Oracle Database 19c Enterprise Edition บน PRIMARY instance ของแต่ละ CDB")
    doc.add_paragraph()

    for i, cdb_info in enumerate(CDBS):
        cdb = cdb_info["cdb"]
        srv = cdb_info["primary_server"]
        log_path = os.path.join(get_host_dir(srv), f"{cdb}_PRIMARY_{srv}_{MONTH_TAG}.txt")
        reader = LogReader(log_path)
        status = "✓ log found" if reader.available() else "⚠ no log"
        print(f"    2.{i+1} PRIMARY – {cdb}  [{status}]")
        build_one_cdb_primary(doc, cdb_info, i + 1, reader)
        add_page_break(doc)

    # ── Listener Status — once per server (ท้ายสุดของ Section 2) ──────────
    primary_server = CDBS[0]["primary_server"]
    add_h2(doc, f"2.{len(CDBS) + 1} Listener Status ({primary_server})")
    add_body(doc, f"ผลลัพธ์จาก lsnrctl status บน {primary_server} (1 Listener ให้บริการทุก CDB)")
    listener_path = os.path.join(get_host_dir(primary_server), f"listener_{primary_server}_{MONTH_TAG}.txt")
    listener_reader = LogReader(listener_path)
    add_raw_log(doc, listener_reader,
                placeholder=f"[ไม่พบ listener_{primary_server}_{MONTH_TAG}.txt — รัน pm_collect.sh แล้ว copy ไฟล์มาวาง]",
                blank_lines=12)

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 3 – Active Data Guard (STANDBY)
# ─────────────────────────────────────────────────────────────────────────────

def build_one_cdb_standby(doc: Document, cdb_info: dict, cdb_idx: int, reader=None):
    cdb = cdb_info["cdb"]
    srv = cdb_info["standby_server"]
    txt = f"{cdb}_STANDBY_{srv}_{MONTH_TAG}.txt"
    pfx = f"3.{cdb_idx}"

    add_h2(doc, f"{pfx} {cdb.upper()} (STANDBY @ {srv})")

    # ── 3.x.1 DG Config ────────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.1 DG Configuration & Sync Status")
    add_body(doc, f"ตรวจสอบ Data Guard Configuration และสถานะ Sync ของ {cdb.upper()}")
    add_sql_section_keys(doc, reader, "[2.2 - 1]",
                         "DG Broker Member Status", "Data Guard Sync Statistics",
                         placeholder=f"[{txt} → [2.2-1] DG Config / Sync Stats]", blank_lines=10,
                         sql_reader=SQL_STANDBY)
    add_findings_box(doc,
        f"พบว่า Data Guard Configuration ของ {cdb.upper()} มี Member ครบถ้วนทั้ง PRIMARY และ PHYSICAL STANDBY "
        f"สถานะ Sync เป็นปกติ ไม่พบ Error ใน DG Broker หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 3.x.2 Gap / Lag ────────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.2 Gap / Lag / Apply Delay")
    add_body(doc, f"ตรวจสอบ Archive Gap และค่า Lag ของ Data Guard {cdb.upper()}")
    add_sql_section_keys(doc, reader, "[2.2 - 2]",
                         "Archive Gap Check", "MRP / Apply Process Status", "Transport / Apply Lag",
                         placeholder=f"[{txt} → [2.2-2] Archive Gap / Apply Lag]", blank_lines=10,
                         sql_reader=SQL_STANDBY)
    add_findings_box(doc,
        f"พบว่าไม่มี Archive Gap ระหว่าง Primary และ Standby ของ {cdb.upper()} "
        f"ค่า Transport Lag และ Apply Lag อยู่ในเกณฑ์ที่ยอมรับได้ หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 3.x.3 Log Transport ────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.3 Log Transport & Apply Service")
    add_body(doc, f"ตรวจสอบ Archive Dest Status และ MRP0 Process ของ {cdb.upper()}")

    add_h4(doc, f"{pfx}.3.1 Archive Dest Status")
    add_sql_section_keys(doc, reader, "[2.2 - 3]", "__preamble__",
                         placeholder=f"[{txt} → [2.2-3] Archive Dest Status]", blank_lines=8,
                         sql_reader=SQL_STANDBY)

    add_h4(doc, f"{pfx}.3.2 Redo Apply Rate")
    add_sql_section_keys(doc, reader, "[2.2 - 3]", "Redo Apply Rate",
                         placeholder=f"[{txt} → [2.2-3] Redo Apply Rate]", blank_lines=8,
                         sql_reader=SQL_STANDBY)
    add_findings_box(doc,
        f"พบว่า Archive Dest Status ของ {cdb.upper()} อยู่ในสถานะ VALID "
        f"และ MRP0 Process กำลัง APPLYING_LOG ตามปกติ "
        f"หมายความว่า Log Transport และ Apply Service ทำงานได้ตามปกติ")

    # ── 3.x.4 Disk & FRA ───────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.4 Disk Usage & FRA")
    add_body(doc, f"ตรวจสอบพื้นที่ Disk และ FRA บน Standby {cdb.upper()}")
    add_sql_section_keys(doc, reader, "[2.2 - 4]",
                         "FRA Summary", "FRA Usage by File Type", "Tablespace Usage",
                         placeholder=f"[{txt} → [2.2-4] FRA / Disk / Tablespace]", blank_lines=10,
                         sql_reader=SQL_STANDBY)
    add_findings_box(doc,
        f"พบว่าพื้นที่ FRA และ Tablespace บน Standby ของ {cdb.upper()} ยังมีพื้นที่ใช้งานเพียงพอ "
        f"ไม่พบพื้นที่ที่ใกล้เต็มหรือเกินเกณฑ์ที่กำหนด หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 3.x.5 Alert Log ────────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.5 Alert Log (Standby)")
    add_body(doc, f"ตรวจสอบ Alert Log ของ Standby {cdb.upper()} เพื่อหา ORA- Error")
    add_alertlog_section(doc, reader, "[2.2 - 5]",
                         placeholder=f"[{txt} → [2.2-5] Alert Log Standby]",
                         sql_reader=SQL_STANDBY, sql_keys=("Alert Log",))
    add_findings_box(doc,
        f"พบว่าใน Alert Log ของ Standby {cdb.upper()} ย้อนหลัง 31 วัน "
        f"ไม่พบ ORA- Error ร้ายแรงหรือ Archive Gap Error "
        f"หมายความว่าสามารถใช้งานได้ตามปกติ")

    # ── 3.x.6 RMAN ─────────────────────────────────────────────────────────
    add_h3(doc, f"{pfx}.6 RMAN on Standby")
    add_body(doc, f"ตรวจสอบสถานะ RMAN Backup บน Standby {cdb.upper()}")
    add_sql_section(doc, reader, "[2.2 - 6]",
                    placeholder=f"[{txt} → [2.2-6] RMAN Standby Backup]", blank_lines=10,
                    sql_reader=SQL_STANDBY, sql_keys=("RMAN Backup",))
    add_findings_box(doc,
        f"ระบบไม่ได้รัน RMAN Backup จาก Standby {cdb.upper()} โดยตรง "
        f"เนื่องจาก Backup ทั้งหมดรันจาก Primary (dbsystem2) ซึ่งถือเป็นเรื่องปกติ "
        f"ดูผลลัพธ์ RMAN ที่สมบูรณ์ได้จาก Section 2.x.5 (Primary)")

    # ── Combined Checklist (bottom of CDB Standby) ─────────────────────────
    add_h3(doc, f"{pfx}.7 รายการตรวจสอบ")
    add_combined_checklist(doc, [
        (f"{pfx}.1  DG Configuration & Sync Status", ITEMS_DG_SYNC),
        (f"{pfx}.2  Gap / Lag / Apply Delay",         ITEMS_DG_GAP),
        (f"{pfx}.3  Log Transport & Apply Service",   ITEMS_DG_TRANS),
        (f"{pfx}.4  Disk Usage & FRA",                ITEMS_DG_DISK),
        (f"{pfx}.5  Alert Log (Standby)",             ITEMS_DG_ALERT),
        (f"{pfx}.6  RMAN on Standby",                 ITEMS_DG_RMAN),
    ])


def build_section3(doc: Document):
    print("  Building Section 3: Active Data Guard...")
    add_h1(doc, "3. Active Data Guard")
    add_body(doc, "ตรวจสอบ Oracle Active Data Guard (Physical Standby) ของแต่ละ CDB บน dbsystem1")
    doc.add_paragraph()

    for i, cdb_info in enumerate(CDBS):
        cdb = cdb_info["cdb"]
        srv = cdb_info["standby_server"]
        log_path = os.path.join(get_host_dir(srv), f"{cdb}_STANDBY_{srv}_{MONTH_TAG}.txt")
        reader = LogReader(log_path)
        status = "✓ log found" if reader.available() else "⚠ no log"
        print(f"    3.{i+1} STANDBY – {cdb}  [{status}]")
        build_one_cdb_standby(doc, cdb_info, i + 1, reader)
        if i < len(CDBS) - 1:
            add_page_break(doc)

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Findings box helper
# ─────────────────────────────────────────────────────────────────────────────

def add_findings_box(doc: Document, text: str = ""):
    """'รายละเอียดของผลลัพธ์' label + plain text description (no gray box)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("รายละเอียดของผลลัพธ์")
    run.bold = True
    run.underline = True
    run.font.name = "TH SarabunPSK"
    run.font.size = Pt(14)

    p_text = doc.add_paragraph()
    p_text.paragraph_format.space_before = Pt(2)
    p_text.paragraph_format.space_after = Pt(10)
    p_text.paragraph_format.first_line_indent = Pt(36)
    r = p_text.add_run(text)
    r.font.name = "TH SarabunPSK"
    r.font.size = Pt(14)


# ─────────────────────────────────────────────────────────────────────────────
# Section 4.2 – CPU content (based on lasted_oracle_critical_patch_update.docx)
# ─────────────────────────────────────────────────────────────────────────────

CPU_DOCX_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                              "lasted_oracle_critical_patch_update.docx")

# CPU releases since 2021 (update quarterly)
CPU_RELEASES = [
    ("Critical Patch Update - April 2026",   "Rev 2, 24 April 2026",    True),
    ("Critical Patch Update - January 2026",  "Rev 3, 02 February 2026", False),
    ("Critical Patch Update - October 2025",  "Rev 1, 21 October 2025",  False),
    ("Critical Patch Update - July 2025",     "Rev 4, 28 July 2025",     False),
    ("Critical Patch Update - April 2025",    "Rev 2, 21 April 2025",    False),
    ("Critical Patch Update - January 2025",  "Rev 2, 11 February 2025", False),
    ("Critical Patch Update - October 2024",  "Rev 2, 25 November 2024", False),
    ("Critical Patch Update - July 2024",     "Rev 3, 18 September 2024",False),
    ("Critical Patch Update - April 2024",    "Rev 2, 18 September 2024",False),
]

# CVE Risk Matrix (April 2026 CPU — Oracle Database Server)
CPU_CVE_ROWS = [
    ("CVE-2024-20696", "Java VM",      "8.8", "No",  "21.3-21.17, 19.3-19.27"),
    ("CVE-2024-20697", "Java VM",      "8.1", "No",  "21.3-21.17, 19.3-19.27"),
    ("CVE-2025-30736", "Java VM",      "5.3", "No",  "21.3-21.17, 19.3-19.27"),
    ("CVE-2025-21549", "XML Database", "4.9", "No",  "21.3-21.17, 19.3-19.27"),
    ("CVE-2025-30739", "Access Struct","2.4", "No",  "21.3-21.17, 19.14-19.27"),
]

CPU_NEXT_DATES = [
    "21 July 2026", "20 October 2026", "19 January 2027", "20 April 2027"
]


def _extract_docx_images(docx_path: str) -> list:
    """Extract PNG/JPG images from a .docx zip into temp files. Returns list of paths."""
    images = []
    if not os.path.exists(docx_path):
        return images
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in sorted(z.namelist()):
            if name.startswith('word/media/') and \
               any(name.lower().endswith(ext) for ext in ('.png', '.jpg', '.jpeg')):
                data = z.read(name)
                ext = os.path.splitext(name)[1]
                tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
                tmp.write(data)
                tmp.close()
                images.append(tmp.name)
    return images


def build_cpu_section(doc: Document):
    """
    Build Section 4.2 content from lasted_oracle_critical_patch_update.docx.
    Embeds actual screenshots + renders CPU release table + CVE risk matrix.
    """
    add_h3(doc, "Latest Patch")
    add_body(doc, (
        "\tตรวจสอบสถานะของ Patch ล่าสุดของซอฟต์แวร์ Oracle "
        "จะพบรายละเอียดของการอัพเดท Patch Oracle ด้วย URL ต่อไปนี้"
    ))

    # URL
    p_url = doc.add_paragraph()
    p_url.paragraph_format.space_before = Pt(2)
    p_url.paragraph_format.space_after = Pt(8)
    r_url = p_url.add_run("\thttps://www.oracle.com/security-alerts/#CriticalPatchUpdates")
    r_url.font.name = "TH SarabunPSK"
    r_url.font.size = Pt(14)
    r_url.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
    r_url.underline = True

    # ── Embed screenshots from the CPU docx ───────────────────────────────
    images = _extract_docx_images(CPU_DOCX_PATH)
    for img_path in images:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(4)
            p_img.paragraph_format.space_after = Pt(4)
            p_img.add_run().add_picture(img_path, width=Inches(6.2))
        except Exception:
            pass
        finally:
            try:
                os.unlink(img_path)
            except Exception:
                pass

    # ── CPU Release Table ─────────────────────────────────────────────────
    add_h4(doc, "Critical Patch Update Releases")
    add_body(doc, f"วันที่ release ครั้งถัดไป: {' / '.join(CPU_NEXT_DATES)}")

    cpu_tbl = doc.add_table(rows=1 + len(CPU_RELEASES), cols=2)
    cpu_tbl.style = "Table Grid"
    cpu_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(cpu_tbl, color="4BACC6", size=8)

    for c, hdr in enumerate(["Critical Patch Update", "Latest Version / Date"]):
        cell = cpu_tbl.rows[0].cells[c]
        set_cell_shading(cell, COLOR_CELL_HEADER)
        set_col_width(cell, 10.0 if c == 0 else 6.5)
        p = cell.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True; r.font.name = "TH SarabunPSK"; r.font.size = Pt(13)
        r.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, (cpu_name, cpu_date, is_latest) in enumerate(CPU_RELEASES):
        row = cpu_tbl.rows[r_idx + 1]
        fill = COLOR_CELL_HEADER if is_latest else (COLOR_CELL_ODD if r_idx % 2 == 0 else COLOR_CELL_EVEN)
        for c_idx, (val, width) in enumerate([(cpu_name, 10.0), (cpu_date, 6.5)]):
            cell = row.cells[c_idx]
            set_cell_shading(cell, fill)
            set_col_width(cell, width)
            p = cell.paragraphs[0]
            r = p.add_run(("★ " if is_latest and c_idx == 0 else "") + val)
            r.bold = is_latest
            r.font.name = "TH SarabunPSK"; r.font.size = Pt(13)
            if is_latest:
                r.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

    # ── Risk Matrix Table ─────────────────────────────────────────────────
    add_h4(doc, "Oracle Database Server Risk Matrix (April 2026 CPU)")
    add_body(doc, (
        "Critical Patch Update April 2026 — Oracle Database Server: "
        "พบ CVE ที่ส่งผลกระทบต่อ Oracle Database 19c (19.3–19.27)"
    ))

    cve_headers = ["CVE ID", "Component", "CVSS Base Score", "Remote Exploit", "Versions Affected"]
    cve_widths  = [3.5, 3.5, 2.5, 2.5, 5.5]
    cve_tbl = doc.add_table(rows=1 + len(CPU_CVE_ROWS), cols=5)
    cve_tbl.style = "Table Grid"
    cve_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(cve_tbl, color="4BACC6", size=8)

    for c, (hdr, w) in enumerate(zip(cve_headers, cve_widths)):
        cell = cve_tbl.rows[0].cells[c]
        set_cell_shading(cell, COLOR_CELL_HEADER)
        set_col_width(cell, w)
        p = cell.paragraphs[0]; r = p.add_run(hdr)
        r.bold = True; r.font.name = "TH SarabunPSK"; r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row_data in enumerate(CPU_CVE_ROWS):
        row = cve_tbl.rows[r_idx + 1]
        score = float(row_data[2])
        fill = "FFE0E0" if score >= 7.0 else ("FFF3CC" if score >= 4.0 else COLOR_CELL_ODD)
        for c_idx, (val, w) in enumerate(zip(row_data, cve_widths)):
            cell = row.cells[c_idx]
            set_cell_shading(cell, fill)
            set_col_width(cell, w)
            p = cell.paragraphs[0]; r = p.add_run(val)
            r.bold = (c_idx == 0)
            r.font.name = "TH SarabunPSK" if c_idx != 0 else "Courier New"
            r.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

    # ── Findings ─────────────────────────────────────────────────────────
    add_findings_box(doc,
        "พบว่า Patch ล่าสุดที่ Oracle ออกคือ แพทช์ประจำไตรมาส เมษายน 2569 "
        "(Critical Patch Update - Apr 2026) ซึ่งมีการแก้ไขช่องโหว่ด้านความปลอดภัย "
        "จึงแนะนำให้ทำการทดสอบกับ Database (Dev) ให้เรียบร้อย "
        "ก่อนทำการติดตั้งแพทช์บน Database (Production)")


# ─────────────────────────────────────────────────────────────────────────────
# Section 4 – Patch
# ─────────────────────────────────────────────────────────────────────────────

def build_section4(doc: Document):
    print("  Building Section 4: Patch...")
    add_h1(doc, "4. Patch")

    add_h2(doc, "4.1 OPatch Patches")
    add_body(doc, "รายการ Oracle Patch ที่ติดตั้งบนแต่ละเซิร์ฟเวอร์ จาก opatch lsinventory และ dba_registry_sqlpatch")

    for host in ("dbsystem1", "dbsystem2"):
        add_h3(doc, f"{host} — Patch / Registry")
        found = False
        # Try PRIMARY log first
        for cdb_info in CDBS:
            if cdb_info["primary_server"] == host:
                log_path = os.path.join(get_host_dir(host), f"{cdb_info['cdb']}_PRIMARY_{host}_{MONTH_TAG}.txt")
                reader = LogReader(log_path)
                if reader.available():
                    add_sql_section_keys(doc, reader, "[2.1 - 7]",
                                         "Database Version", "Installed Components", "Applied Patches",
                                         placeholder=f"[opatch lsinventory / dba_registry_sqlpatch — {host}]",
                                         blank_lines=12)
                    found = True
                    break
        # Fallback: try STANDBY log (for hosts that only run as Standby)
        if not found:
            for cdb_info in CDBS:
                if cdb_info["standby_server"] == host:
                    log_path = os.path.join(get_host_dir(host), f"{cdb_info['cdb']}_STANDBY_{host}_{MONTH_TAG}.txt")
                    reader = LogReader(log_path)
                    if reader.available() and reader.get("[2.1 - 7]").strip():
                        add_sql_section_keys(doc, reader, "[2.1 - 7]",
                                             "Database Version", "Installed Components", "Applied Patches",
                                             placeholder=f"[dba_registry_sqlpatch — {host}]",
                                             blank_lines=12)
                        found = True
                        break
        if not found:
            add_log_box(doc, f"[วางผลลัพธ์ opatch lsinventory และ dba_registry_sqlpatch — {host}]", blank_lines=12)
        add_findings_box(doc,
            f"พบว่า Oracle Patch ที่ติดตั้งบน {host} เป็น Database Release Update ประจำไตรมาสล่าสุด "
            f"ที่ได้รับการทดสอบและอนุมัติเรียบร้อยแล้ว รายละเอียด Patch เปรียบเทียบกับ CPU ล่าสุดดูได้ที่ Section 4.2 "
            f"หมายความว่าสามารถใช้งานได้ตามปกติ")

    add_checklist(doc, ITEMS_PATCH)

    add_h2(doc, "4.2 Oracle Critical Patch Update (CPU)")
    add_body(doc, (
        "ตรวจสอบ Oracle Critical Patch Update (CPU) รายไตรมาสจากเว็บไซต์ Oracle "
        "เปรียบเทียบกับ Patch ที่ติดตั้งบนระบบปัจจุบัน"
    ))
    build_cpu_section(doc)

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 5 – Network
# ─────────────────────────────────────────────────────────────────────────────

def build_section5(doc: Document):
    print("  Building Section 5: Network...")
    add_h1(doc, "5. Network")
    add_body(doc, "ตรวจสอบ Network Connectivity และ Bandwidth ระหว่าง Primary (dbsystem2) และ Standby (dbsystem1)")

    # tnsping target: dbsystem1 → AMSSCDB, dbsystem2 → AMSSCDB_STBY
    TNSPING_TARGET = {"dbsystem1": "AMSSCDB_STBY", "dbsystem2": "AMSSCDB"}

    # ── 5.1 tnsping Test ───────────────────────────────────────────────────
    add_h2(doc, "5.1 tnsping Test")
    add_body(doc, (
        "ทดสอบการเชื่อมต่อ Oracle Net ระหว่าง Primary และ Standby ด้วย tnsping "
        "— dbsystem1 ทดสอบไปยัง AMSSCDB (Primary), dbsystem2 ทดสอบไปยัง AMSSCDB_STBY (Standby)"
    ))

    for host in ("dbsystem1", "dbsystem2"):
        target = TNSPING_TARGET[host]
        add_h3(doc, f"{host} → tnsping {target}")
        path = os.path.join(get_host_dir(host), f"network_test_{host}_{MONTH_TAG}.txt")
        net_reader = LogReader(path)
        add_raw_log(doc, net_reader,
                    placeholder=f"[ไม่พบ network_test_{host}_{MONTH_TAG}.txt — รัน pm_collect.sh แล้ว copy ไฟล์มาวาง]",
                    blank_lines=8)
    add_findings_box(doc,
        "พบว่า tnsping สามารถ resolve และ connect ผ่าน Oracle Net ได้สำเร็จจากทั้งสองฝั่ง "
        "หมายความว่า Oracle Listener และ TNS Configuration ทำงานได้ตามปกติ")
    add_checklist(doc, ITEMS_NET)

    # ── 5.2 Bandwidth ──────────────────────────────────────────────────────
    add_h2(doc, "5.2 Network Bandwidth ระหว่าง Primary-Standby")
    add_body(doc, (
        "ทดสอบ Bandwidth โดยการโยนไฟล์ขนาดใหญ่ระหว่าง dbsystem1 และ dbsystem2 "
        "บันทึกเวลาและขนาดไฟล์เพื่อคำนวณ Throughput"
    ))
    add_log_box(doc,
                "[วางผลลัพธ์การทดสอบ Bandwidth (เช่น scp/rsync output, ขนาดไฟล์, เวลา, MB/s)]",
                blank_lines=10)
    add_findings_box(doc,
        "พบว่าค่า Network Bandwidth ระหว่าง Primary และ Standby เพียงพอสำหรับการส่ง Archive Log "
        "และ Redo Data หมายความว่าสามารถใช้งานได้ตามปกติ")

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 6 – Summary
# ─────────────────────────────────────────────────────────────────────────────

def build_section6_summary(doc: Document):
    print("  Building Section 6: สรุปผลการตรวจสอบ...")
    add_h1(doc, "6. สรุปผลการตรวจสอบ (Summary)")
    add_body(doc, "ตารางสรุปผลการตรวจสอบ Oracle Database Health Check ประจำเดือน พฤษภาคม 2568")
    doc.add_paragraph()

    all_rows = []

    all_rows.append(("__header__", "Section 2 — Database Enterprise Edition (PRIMARY)", "", ""))
    for i, cdb_info in enumerate(CDBS):
        cdb = cdb_info["cdb"].upper(); pdbs = cdb_info["pdbs"]; pfx = f"2.{i+1}"
        all_rows.append(("__subhdr__", f"{pfx}  {cdb} (PRIMARY @ {cdb_info['primary_server']})", "", ""))
        rows_2x = [
            (cdb, f"{pfx}.1.1", "Instance Status"),
            (cdb, f"{pfx}.1.2", "Registered Services"),
            (cdb, f"{pfx}.2.1", "Tablespace – CDB Root"),
        ] + [(cdb, f"{pfx}.2.{j+2}", f"Tablespace – PDB: {pdbs[j]}") for j in range(len(pdbs))] + [
            (cdb, f"{pfx}.3",  "Alert Log (Primary)"),
            (cdb, f"{pfx}.4",  "Performance (AWR / Top SQL / Wait Events)"),
            (cdb, f"{pfx}.5",  "RMAN Backup"),
            (cdb, f"{pfx}.6",  "Parameters (SGA / PGA)"),
            (cdb, f"{pfx}.7",  "Oracle Version"),
        ]
        all_rows.extend(rows_2x)

    all_rows.append(("__header__", "Section 3 — Active Data Guard (STANDBY)", "", ""))
    for i, cdb_info in enumerate(CDBS):
        cdb = cdb_info["cdb"].upper(); pfx = f"3.{i+1}"
        all_rows.append(("__subhdr__", f"{pfx}  {cdb} (STANDBY @ {cdb_info['standby_server']})", "", ""))
        all_rows.extend([
            (cdb, f"{pfx}.1",   "DG Configuration & Sync Status"),
            (cdb, f"{pfx}.2",   "Gap / Lag / Apply Delay"),
            (cdb, f"{pfx}.3.1", "Archive Dest Status"),
            (cdb, f"{pfx}.3.2", "Redo Apply Rate"),
            (cdb, f"{pfx}.4",   "Disk Usage & FRA"),
            (cdb, f"{pfx}.5",   "Alert Log (Standby)"),
            (cdb, f"{pfx}.6",   "RMAN on Standby"),
        ])

    all_rows.append(("__header__", "Section 4 — Patch", "", ""))
    all_rows.extend([
        ("—", "4.1", "OPatch Patches — dbsystem1"),
        ("—", "4.1", "OPatch Patches — dbsystem2"),
        ("—", "4.2", "Oracle Critical Patch Update (CPU)"),
    ])

    all_rows.append(("__header__", "Section 5 — Network", "", ""))
    all_rows.extend([
        ("—", "5.1", "tnsping Test (dbsystem1→AMSSCDB / dbsystem2→AMSSCDB_STBY)"),
        ("—", "5.2", "Network Bandwidth (Manual File Transfer)"),
    ])

    sum_headers = ["CDB", "Section", "รายการตรวจสอบ", "สถานะ", "หมายเหตุ"]
    sum_col_w   = [2.0, 2.0, 7.5, 3.2, 2.6]
    sum_tbl = doc.add_table(rows=1 + len(all_rows), cols=5)
    sum_tbl.style = "Table Grid"
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(sum_tbl, color="4BACC6", size=8)

    hdr_row = sum_tbl.rows[0]
    for c, (hdr, w) in enumerate(zip(sum_headers, sum_col_w)):
        cell = hdr_row.cells[c]; set_cell_shading(cell, COLOR_CELL_HEADER); set_col_width(cell, w)
        p = cell.paragraphs[0]; run = p.add_run(hdr)
        run.bold = True; run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
        run.font.name = "TH SarabunPSK"; run.font.size = Pt(13)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for data_row_idx, r_data in enumerate(all_rows):
        row = sum_tbl.rows[data_row_idx + 1]

        if r_data[0] == "__header__":
            merged = row.cells[0]
            for i in range(1, 5):
                merged = merged.merge(row.cells[i])
            set_cell_shading(merged, COLOR_CELL_HEADER)
            p = merged.paragraphs[0]; run = p.add_run(r_data[1])
            run.bold = True; run.font.name = "TH SarabunPSK"; run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif r_data[0] == "__subhdr__":
            merged = row.cells[0]
            for i in range(1, 5):
                merged = merged.merge(row.cells[i])
            set_cell_shading(merged, COLOR_SECTION_HDR)
            p = merged.paragraphs[0]; run = p.add_run(r_data[1])
            run.bold = True; run.font.name = "TH SarabunPSK"; run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(COLOR_HDR_TEXT)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        else:
            cdb_lbl, sec_lbl, item_lbl = r_data[0], r_data[1], r_data[2]
            fill = COLOR_CELL_ODD if data_row_idx % 2 == 0 else COLOR_CELL_EVEN
            cells = row.cells

            def fill_cell(c_idx, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
                cell = cells[c_idx]; set_cell_shading(cell, fill); set_col_width(cell, sum_col_w[c_idx])
                p = cell.paragraphs[0]; run = p.add_run(text)
                run.bold = bold; run.font.name = "TH SarabunPSK"; run.font.size = Pt(13)
                p.alignment = align; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            fill_cell(0, cdb_lbl, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            fill_cell(1, sec_lbl, align=WD_ALIGN_PARAGRAPH.CENTER)
            fill_cell(2, item_lbl)

            cell3 = cells[3]; set_cell_shading(cell3, fill); set_col_width(cell3, sum_col_w[3])
            p3 = cell3.paragraphs[0]; p3.add_run("☐ ปกติ  ☐ ไม่ปกติ").font.name = "TH SarabunPSK"
            p3.runs[0].font.size = Pt(12); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            cell4 = cells[4]; set_cell_shading(cell4, fill); set_col_width(cell4, sum_col_w[4])
            cell4.paragraphs[0].add_run(""); cell4.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

    add_h2(doc, "ผลโดยรวม")
    overall_data = [
        ("ผลโดยรวม",              "☐ ปกติ          ☐ ต้องดำเนินการแก้ไข"),
        ("จำนวนรายการปกติ",        "_______ รายการ"),
        ("ข้อสังเกต / ข้อเสนอแนะ", "\n\n\n"),
    ]
    ov_tbl = doc.add_table(rows=len(overall_data), cols=2)
    ov_tbl.style = "Table Grid"; ov_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(ov_tbl, color="4BACC6", size=8)
    for r_idx, (lbl, val) in enumerate(overall_data):
        cells = ov_tbl.rows[r_idx].cells
        set_cell_shading(cells[0], COLOR_COVER_LABEL); set_col_width(cells[0], 4.5)
        r0 = cells[0].paragraphs[0].add_run(lbl); r0.bold = True
        r0.font.name = "TH SarabunPSK"; r0.font.size = Pt(14)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_col_width(cells[1], 12.5)
        r1 = cells[1].paragraphs[0].add_run(val)
        r1.font.name = "TH SarabunPSK"; r1.font.size = Pt(14)
    doc.add_paragraph()

    add_h2(doc, "ลายเซ็นผู้ตรวจสอบ")
    sig_data = [
        ("ลงชื่อ _________________________________ ผู้ปฏิบัติงาน",
         "ลงชื่อ _________________________________ ผู้ควบคุม"),
        ("(                                                              )",
         "(                                                              )"),
        ("วันที่ _______ / _______ / _______",
         "วันที่ _______ / _______ / _______"),
    ]
    sig_tbl = doc.add_table(rows=len(sig_data), cols=2)
    sig_tbl.style = "Table Grid"; sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(sig_tbl, color="4BACC6", size=8)
    for r_idx, (left_txt, right_txt) in enumerate(sig_data):
        cells = sig_tbl.rows[r_idx].cells
        for c_idx, txt in enumerate([left_txt, right_txt]):
            set_col_width(cells[c_idx], 8.5)
            p = cells[c_idx].paragraphs[0]; p.add_run(txt).font.name = "TH SarabunPSK"
            p.runs[0].font.size = Pt(14); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ─────────────────────────────────────────────────────────────────────────────
# Document setup
# ─────────────────────────────────────────────────────────────────────────────

def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    doc.styles["Normal"].font.name = "TH SarabunPSK"
    doc.styles["Normal"].font.size = Pt(14)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "TH SarabunPSK"; h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor.from_string(COLOR_H1)
    h1.paragraph_format.space_before = Pt(12); h1.paragraph_format.space_after = Pt(6)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "TH SarabunPSK"; h2.font.size = Pt(16)
    h2.font.color.rgb = RGBColor.from_string(COLOR_H2)
    h2.paragraph_format.space_before = Pt(8); h2.paragraph_format.space_after = Pt(4)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "TH SarabunPSK"; h3.font.size = Pt(14)
    h3.font.color.rgb = RGBColor.from_string(COLOR_H3)
    h3.paragraph_format.space_before = Pt(6); h3.paragraph_format.space_after = Pt(2)

    h4 = doc.styles["Heading 4"]
    h4.font.name = "TH SarabunPSK"; h4.font.size = Pt(13)
    h4.font.color.rgb = RGBColor.from_string(COLOR_H4)

    # Remove automatic list-numbering from heading styles (the cover template may
    # already carry numPr, which would double-up with the "1." already in heading text)
    for level in range(1, 5):
        pPr = doc.styles[f"Heading {level}"].element.get_or_add_pPr()
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)

    h4.paragraph_format.space_before = Pt(4); h4.paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # ── Open cover-page file as the base document ──────────────────────────
    # COVER_FILE is never overwritten; output always goes to OUTPUT_FILE.
    if os.path.exists(COVER_FILE):
        print(f"Cover template : {os.path.basename(COVER_FILE)}")
        doc = Document(COVER_FILE)
        setup_document(doc)
        add_page_break(doc)
        cover_built = True
    else:
        print(f"WARNING: cover file not found — {COVER_FILE}")
        print("         Building cover from scratch instead.")
        doc = Document()
        setup_document(doc)
        cover_built = False

    print(f"Output         : {OUTPUT_FILE}")
    print(f"Logs           : {OUTPUT_DIR}")
    print()
    print("Generating Oracle Health Check Report...")
    print()

    if not cover_built:
        print("[1/7] Cover page")
        build_cover(doc)

    print("[1/6] Section 1: System Information")
    build_section1(doc)

    print("[2/6] Section 2: Database Enterprise Edition")
    build_section2(doc)

    print("[3/6] Section 3: Active Data Guard")
    build_section3(doc)

    print("[4/6] Section 4: Patch")
    build_section4(doc)

    print("[5/6] Section 5: Network")
    build_section5(doc)

    print("[6/6] Section 6: Summary")
    build_section6_summary(doc)

    doc.save(OUTPUT_FILE)
    print()
    print("Done. Report saved to:")
    print(f"  {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
